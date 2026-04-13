from fastapi import APIRouter, BackgroundTasks, HTTPException, WebSocket, WebSocketDisconnect, File, Form, UploadFile
from pydantic import BaseModel, Field
import uuid
import io
import PyPDF2
import docx
from src.config.database import get_db
from src.models.schema import ExecutionDocument, PRDDocument, PipelineStatus
from src.control_plane.pipeline import run_pipeline
from src.api.websocket import manager

router = APIRouter()

class PRDSubmitRequest(BaseModel):
    title: str = Field(..., min_length=1, max_length=500)
    content: str = Field(..., min_length=10)
    metadata: dict = None

class PRDSubmitResponse(BaseModel):
    execution_id: str
    prd_id: str
    status: str
    message: str

@router.post("/api/v1/prd", response_model=PRDSubmitResponse)
async def submit_prd(
    background_tasks: BackgroundTasks,
    title: str = Form("Auto-Generated"),
    content: str = Form(""),
    file: UploadFile = File(None)
):
    db = get_db()
    
    extracted_text = ""
    if file:
        file_bytes = await file.read()
        if file.filename.endswith(".pdf"):
            pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_bytes))
            for page in pdf_reader.pages:
                extracted_text += page.extract_text() + "\n"
        elif file.filename.endswith(".docx") or file.filename.endswith(".doc"):
            doc = docx.Document(io.BytesIO(file_bytes))
            extracted_text = "\n".join([para.text for para in doc.paragraphs])
        else:
            extracted_text = file_bytes.decode('utf-8', errors='ignore')
            
    final_content = f"{content}\n\n{extracted_text}".strip()
    if not final_content:
        raise HTTPException(status_code=400, detail="Must provide text content or a file")
        
    # 1. Insert PRD
    prd = PRDDocument(raw_text=final_content)
    await db["prds"].insert_one(prd.model_dump(mode='json'))
    
    # 2. Setup Execution
    execution = ExecutionDocument(prd_id=prd.id, status=PipelineStatus.PENDING)
    await db["executions"].insert_one(execution.model_dump(mode='json'))
    
    # 3. Queue Pipeline
    background_tasks.add_task(run_pipeline, prd.id, execution.id, final_content)
    
    return PRDSubmitResponse(
        execution_id=execution.id,
        prd_id=prd.id,
        status="accepted",
        message="PRD submitted successfully. Workflow execution started."
    )

@router.websocket("/ws")
async def websocket_endpoint(websocket: WebSocket):
    await manager.connect(websocket)
    try:
        while True:
            # We just keep connection alive, frontend doesn't send data usually
            data = await websocket.receive_text()
            await manager.broadcast({"message": data})
    except WebSocketDisconnect:
        manager.disconnect(websocket)
